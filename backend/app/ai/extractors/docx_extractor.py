import docx
from typing import Dict, Any, List

class DOCXExtractor:
    def extract(self, file_path: str) -> Dict[str, Any]:
        """Extracts text, headings, and table data from DOCX files."""
        paragraphs: List[str] = []
        tables_data: List[str] = []

        try:
            doc = docx.Document(file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())

            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_rows.append(row_text)
                if table_rows:
                    tables_data.append("\n".join(table_rows))
        except Exception as e:
            paragraphs.append(f"DOCX extraction note: {str(e)}")

        full_text = "\n\n".join(paragraphs)
        if tables_data:
            full_text += "\n\n--- EXTRACTED TABLES ---\n" + "\n\n".join(tables_data)

        return {
            "full_text": full_text,
            "page_count": 1 + (len(paragraphs) // 15),
            "paragraphs_count": len(paragraphs),
            "has_text": len(full_text.strip()) > 0
        }

docx_extractor = DOCXExtractor()
